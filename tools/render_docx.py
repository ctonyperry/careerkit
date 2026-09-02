"""Render a run's resume.md (and cover-letter.md) to .docx.

Until now this was an undocumented manual step, which meant the one gate that
is measured rather than estimated -- page count -- depended on someone
remembering how the last one was built. The layout here is reverse-engineered
from runs/2026-08-25-edia-forward-deployed-engineer/resume.docx, which is the
package Tony signed off on: 9pt Calibri body, 0.45in top/bottom and 0.5in
left/right margins, tight paragraph spacing, section headings upper-cased.

    python tools/render_docx.py runs/<run-dir>

Requires python-docx. Word itself is only needed for the page measurement in
tools/metrics.py, which reads the file this produces.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

BODY = "Calibri"
INK = RGBColor(0x1A, 0x1A, 0x1A)

# (size, bold, italic, space_before, space_after) in points.
SPEC = {
    "name": (17.0, True, False, None, 1.5),
    "tagline": (10.5, True, False, None, 3.75),
    "contact": (8.5, False, False, None, 4.25),
    "summary": (9.0, False, False, None, 3.5),
    "section": (10.0, True, False, 7.5, 2.75),
    "role": (9.5, True, False, 5.25, 1.0),
    "rolenote": (9.0, False, True, None, 2.25),
    "bullet": (9.0, False, False, None, 2.25),
    "para": (9.0, False, False, None, 3.5),
}

BULLET = "• "
INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")


def _para(doc, kind: str, text: str, *, indent: float | None = None):
    size, bold, italic, before, after = SPEC[kind]
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.0
    if indent is not None:
        pf.left_indent = Inches(indent)
        pf.first_line_indent = Inches(-0.13)
    # **bold** spans render bold; everything else takes the paragraph default.
    pos = 0
    for m in INLINE_BOLD.finditer(text):
        for chunk, forced in ((text[pos : m.start()], None), (m.group(1), True)):
            if not chunk:
                continue
            r = p.add_run(chunk)
            r.font.name, r.font.size = BODY, Pt(size)
            r.font.bold = bold if forced is None else True
            r.font.italic = italic
            r.font.color.rgb = INK
        pos = m.end()
    tail = text[pos:]
    if tail or pos == 0:
        r = p.add_run(tail)
        r.font.name, r.font.size = BODY, Pt(size)
        r.font.bold, r.font.italic = bold, italic
        r.font.color.rgb = INK
    return p


def render_resume(md: str, out: Path) -> Path:
    doc = docx.Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Inches(0.45)
    s.left_margin = s.right_margin = Inches(0.5)

    lines = md.splitlines()
    i, seen_h1 = 0, False
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        i += 1
        if not line:
            continue
        if line.startswith("# "):
            _para(doc, "name", line[2:].strip())
            seen_h1 = True
            # The bolded line right under the name is the tagline.
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines) and lines[i].strip().startswith("**"):
                _para(doc, "tagline", INLINE_BOLD.sub(r"\1", lines[i].strip()))
                i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines) and "·" in lines[i]:
                _para(doc, "contact", lines[i].strip())
                i += 1
        elif line.startswith("## "):
            _para(doc, "section", line[3:].strip().upper())
        elif line.startswith("### "):
            _para(doc, "role", line[4:].strip())
        elif line.startswith("- "):
            _para(doc, "bullet", BULLET + line[2:].strip(), indent=0.13)
        elif line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            _para(doc, "rolenote", line.strip("*").strip())
        else:
            _para(doc, "summary" if seen_h1 and len(doc.paragraphs) < 5 else "para", line)

    doc.save(str(out))
    return out


def render_letter(md: str, out: Path) -> Path:
    doc = docx.Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.9)

    body: list[str] = []
    for block in re.split(r"\n\s*\n", md.strip()):
        block = " ".join(x.strip() for x in block.splitlines() if x.strip())
        if not block or block.startswith("# "):
            continue
        body.append(block)

    # Identity comes from the spine, never from code. The spine's identity
    # block is the same one the resume brief renders from.
    import yaml
    from _paths import SPINE
    ident = yaml.safe_load(SPINE.read_text(encoding="utf-8")).get("identity", {})
    _para(doc, "name", ident.get("resume_header", ""))
    _para(doc, "contact", " · ".join(
        str(ident[k]) for k in ("email", "phone", "location", "linkedin") if ident.get(k)))
    for block in body:
        p = _para(doc, "para", block)
        p.paragraph_format.space_after = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.save(str(out))
    return out


def main(argv: list[str]) -> int:
    if len(argv) != 2:
        print(__doc__)
        return 2
    run = Path(argv[1])
    if not run.is_dir():
        print(f"not a run directory: {run}")
        return 1
    made: list[Path] = []
    locked: list[Path] = []

    def _try(render, src: Path, out: Path) -> None:
        # Tony hand-tunes these in Word, which holds an exclusive lock. That is
        # a normal state, not an error, and it must not take down the render of
        # the other document or look like a crash.
        if not src.exists():
            return
        try:
            made.append(render(src.read_text(encoding="utf-8"), out))
        except PermissionError:
            locked.append(out)

    _try(render_resume, run / "resume.md", run / "resume.docx")
    _try(render_letter, run / "cover-letter.md", run / "cover-letter.docx")

    if not made and not locked:
        print(f"no resume.md or cover-letter.md in {run}")
        return 1
    for m in made:
        print(f"wrote {m}")
    for lock in locked:
        print(f"SKIPPED {lock}: open in another program, markdown is still current")
    return 3 if locked else 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
